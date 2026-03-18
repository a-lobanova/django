from django.shortcuts import render
import os
import pandas as pd
from django.conf import settings
from .forms import UploadFileForm
from docx import Document
import logging

# Создаём логгер для приложения roads
logger = logging.getLogger("roads")


def upload_file(request):
    download_url = None
    error_message = None

    if request.method == "POST":
        form = UploadFileForm(request.POST, request.FILES)
        if form.is_valid():
            excel_file = request.FILES["file"]
            logger.info(f"Загружен файл: {excel_file.name}")

            # --- Сохраняем временно в media/ ---
            os.makedirs(settings.MEDIA_ROOT, exist_ok=True)
            excel_path = os.path.join(settings.MEDIA_ROOT, excel_file.name)
            try:
                with open(excel_path, "wb+") as f:
                    for chunk in excel_file.chunks():
                        f.write(chunk)
                logger.info(f"Файл сохранён во временную директорию: {excel_path}")
            except Exception as e:
                error_message = f"Ошибка сохранения файла: {e}"
                logger.error(error_message)
                return render(
                    request, "roads/upload.html", {"form": form, "error": error_message}
                )

            # --- Чтение Excel ---
            try:
                df = pd.read_excel(excel_path)
                logger.info(
                    f"Файл {excel_file.name} успешно прочитан, строк: {len(df)}"
                )
            except Exception as e:
                error_message = f"Ошибка чтения Excel файла: {e}"
                logger.error(error_message)
                df = None

            if df is not None and df.empty:
                error_message = "Excel файл не содержит данных"
                logger.warning(f"{excel_file.name}: {error_message}")

            # --- Проверка колонок ---
            expected_columns = [
                "№ п/п",
                "Наименование",
                "Значение автомобильной дороги",
                "Категория",
                "Протяженность, км",
            ]
            if df is not None:
                df.columns = df.columns.str.strip()
                if not all(col in df.columns for col in expected_columns):
                    error_message = "Некорректные колонки в Excel файле!"
                    logger.warning(f"{excel_file.name}: {error_message}")

            # --- Проверка чисел ---
            if df is not None and error_message is None:
                try:
                    df["Протяженность, км"] = (
                        df["Протяженность, км"]
                        .astype(str)
                        .str.replace(" ", "")
                        .str.replace(",", ".")
                    )
                    df["Протяженность, км"] = pd.to_numeric(
                        df["Протяженность, км"], errors="coerce"
                    )
                    logger.info(
                        f"Колонка 'Протяженность, км' успешно приведена к числам"
                    )
                except Exception as e:
                    error_message = f"Ошибка преобразования чисел: {e}"
                    logger.error(error_message)

                if error_message is None:
                    if df["Протяженность, км"].isnull().any():
                        error_message = (
                            "Некорректные значения в колонке Протяженность, км"
                        )
                        logger.warning(f"{excel_file.name}: {error_message}")
                    elif (df["Протяженность, км"] < 0).any():
                        error_message = "Протяженность не может быть отрицательной"
                        logger.warning(f"{excel_file.name}: {error_message}")

            # --- Если есть ошибка, возвращаем форму с сообщением ---
            if error_message:
                return render(
                    request, "roads/upload.html", {"form": form, "error": error_message}
                )

            # --- Считаем общую протяженность ---
            total_length = df["Протяженность, км"].sum()
            total_text = (
                f"Общая протяженность автомобильных дорог составляет {total_length} км"
            )
            logger.info(f"{excel_file.name}: {total_text}")

            # --- Генерация Word ---
            try:
                doc = Document()
                doc.add_heading(
                    "Таблица 1 - Перечень и характеристика автомобильных дорог, проходящих по территории муниципального округа",
                    level=2,
                )
                table = doc.add_table(rows=1, cols=len(df.columns))
                table.style = "Table Grid"

                # Заголовки таблицы
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(df.columns):
                    hdr_cells[i].text = str(col)

                # Данные
                for _, row in df.iterrows():
                    row_cells = table.add_row().cells
                    for i, col in enumerate(df.columns):
                        row_cells[i].text = str(row[col])

                doc.add_paragraph(total_text)

                # Сохраняем документ
                filename = f"report_{excel_file.name}.docx"
                output_path = os.path.join(settings.MEDIA_ROOT, filename)
                doc.save(output_path)
                logger.info(f"Отчёт сгенерирован и сохранён: {output_path}")

                download_url = settings.MEDIA_URL + filename

            except Exception as e:
                error_message = f"Ошибка генерации отчёта: {e}"
                logger.error(error_message)
                return render(
                    request, "roads/upload.html", {"form": form, "error": error_message}
                )

    else:
        form = UploadFileForm()

    return render(
        request,
        "roads/upload.html",
        {"form": form, "download_url": download_url},
    )
